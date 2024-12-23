# Reattempt to generate the Word document and provide the correct download link
from docx import Document

# Initialize the document
doc = Document()

# Define formatting constants
font_name = 'Times New Roman'
font_size_regular = 12
font_size_heading = 15

# Create cover page
doc.add_paragraph("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("GALGOTIAS COLLEGE OF ENGINEERING & TECHNOLOGY\nGREATER NOIDA, UTTAR PRADESH, INDIA", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("B.TECH MINI PROJECT REPORT", style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n\nON\n\n", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("HAIR CARE WEBSITE", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n\nSubmitted by\nYour Name\nRoll No: XYZ\n", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

# Add certificate page
doc.add_paragraph("CERTIFICATE", style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("This is to certify that the mini project report titled “Hair Care Website” "
                  "is a bonafide work done by Your Name, Roll No: XYZ under the supervision of Faculty Name "
                  "towards partial fulfillment of the requirements for the award of the degree of Bachelor of Technology "
                  "in Computer Science & Engineering at Galgotias College of Engineering & Technology.")
doc.add_page_break()

# Add acknowledgment page
doc.add_paragraph("ACKNOWLEDGEMENT", style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("I would like to express my sincere gratitude to my guide, Faculty Name, for their invaluable guidance and support "
                  "throughout this project. I would also like to thank the Department of Computer Science & Engineering for providing "
                  "the resources and environment to complete this project successfully. Lastly, I am grateful to my family and friends for their encouragement.")
doc.add_page_break()

# Add abstract page
doc.add_paragraph("ABSTRACT", style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("This project report presents the design and development of a hair care website. "
                  "The website aims to provide users with research-based information on hair health, "
                  "common causes of hair fall, and scientifically supported solutions. It incorporates "
                  "interactive features to assist users in choosing products tailored to their needs.\n"
                  "Keywords: Hair care, hair fall, solutions, website development, user interaction.")
doc.add_page_break()

# Add table of contents (placeholder)
doc.add_paragraph("TABLE OF CONTENTS", style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n1. Introduction\n2. Literature Review\n3. Problem Formulation\n4. Proposed Work\n"
                  "5. System Design\n6. Implementation\n7. Result Analysis\n8. Conclusion, Limitations, and Future Work\n"
                  "9. References").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
doc.add_page_break()

# Add chapters with headings and content placeholders
chapters = [
    "Introduction",
    "Literature Review",
    "Problem Formulation",
    "Proposed Work",
    "System Design",
    "Implementation",
    "Result Analysis",
    "Conclusion, Limitations, and Future Work"
]

content = {
    "Introduction": "The introduction provides an overview of the problem of hair care and the motivation for developing the website.",
    "Literature Review": "This section reviews existing research on hair health, the causes of hair fall, and potential remedies.",
    "Problem Formulation": "Defines the specific challenges and objectives addressed by the website project.",
    "Proposed Work": "Explains the proposed methodology for developing the hair care website.",
    "System Design": "Details the system architecture, including front-end and back-end design.",
    "Implementation": "Describes the tools, technologies, and processes used to build the website.",
    "Result Analysis": "Presents findings, including user feedback and performance metrics.",
    "Conclusion, Limitations, and Future Work": "Summarizes the project outcomes, discusses limitations, and proposes future enhancements."
}

for chapter in chapters:
    doc.add_paragraph(chapter.upper(), style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(content[chapter], style='Normal')

# Save document
file_path = "/mnt/data/Hair_Care_Website_Report_v2.docx"
doc.save(file_path)

file_path
